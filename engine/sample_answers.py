"""
Sample Answer Generator
Generates Band 7+ sample answers using LLM API.
Requires OPENAI_API_KEY environment variable or config file.
"""
import os, json, datetime
import requests

# ---- Configuration ----

_config_loaded = False
_api_key = None
_api_base = None
_model = None


def _load_config():
    """Load LLM configuration from environment variables or config file."""
    global _config_loaded, _api_key, _api_base, _model
    if _config_loaded:
        return
    _api_key = os.environ.get("DEEPSEEK_API_KEY") or os.environ.get("OPENAI_API_KEY") or os.environ.get("LLM_API_KEY")
    _api_base = os.environ.get("DEEPSEEK_BASE_URL") or os.environ.get("OPENAI_BASE_URL") or "https://api.deepseek.com/v1"
    _model = os.environ.get("DEEPSEEK_MODEL") or os.environ.get("LLM_MODEL") or "deepseek-chat"
    if not _api_key:
        cfgs = [
            os.path.join(os.path.dirname(__file__), "..", "llm_config.json"),
            os.path.join(os.path.dirname(__file__), "llm_config.json"),
        ]
        for cfg in cfgs:
            cfg = os.path.abspath(cfg)
            if os.path.exists(cfg):
                try:
                    with open(cfg, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    _api_key = data.get("api_key") or _api_key
                    _api_base = data.get("api_base") or _api_base
                    _model = data.get("model") or _model
                except Exception:
                    pass
                break
    _config_loaded = True


def _call_llm(system_prompt, user_prompt):
    """Call the OpenAI-compatible LLM API and return the response text."""
    _load_config()
    if not _api_key:
        return "[ERROR] LLM API key not configured.\nPlease set OPENAI_API_KEY environment variable or create a llm_config.json file."
    try:
        headers = {
            "Authorization": "Bearer " + _api_key,
            "Content-Type": "application/json",
        }
        url = _api_base.rstrip("/") + "/chat/completions"
        payload = {
            "model": _model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.8,
            "max_tokens": 2048,
        }
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
        if resp.status_code != 200:
            return "[ERROR] LLM API returned status " + str(resp.status_code) + ":\n" + resp.text[:200]
        data = resp.json()
        return data["choices"][0]["message"]["content"].strip()
    except requests.exceptions.Timeout:
        return "[ERROR] LLM API request timed out. Please check your network."
    except requests.exceptions.ConnectionError:
        return "[ERROR] Cannot connect to LLM API. Please check your network and API endpoint."
    except Exception as e:
        return "[ERROR] LLM call failed: " + str(e)


def _build_prompt(part, question_info, user_response, topic):
    """Build the system prompt and user prompt for the LLM."""
    system_prompt = (
        "You are an expert IELTS examiner and English language instructor with over 20 years of experience. "
        "Your task is to generate high-quality sample answers for the IELTS Speaking test that would achieve Band 7.0 or higher. "
        "Analyze the student's response carefully, then create TWO distinct, natural-sounding sample answers "
        "that demonstrate the vocabulary, grammatical structures, coherence, and fluency expected at Band 7+. "
        "Each sample answer should:\n"
        "1. Fully address the question/task\n"
        "2. Use a range of vocabulary including less common items and collocations\n"
        "3. Use a mix of simple and complex sentence structures with good control\n"
        "4. Flow naturally with appropriate discourse markers and linking\n"
        "5. Be the appropriate length for the part (P1: 30-45 seconds, P2: 1-2 minutes, P3: 45-60 seconds)\n"
        "6. Sound like a real person speaking, NOT an essay reader"
    )

    if part == 1 or part == "1":
        part_desc = "Part 1 - Introduction and Interview (short personal questions)"
        if isinstance(question_info, list):
            q_text = "\n".join(question_info)
        else:
            q_text = str(question_info)
    elif part == 2 or part == "2":
        part_desc = "Part 2 - Individual Long Turn (Cue Card, speak for 1-2 minutes)"
        if isinstance(question_info, dict):
            q_text = "Cue Card: " + question_info.get("cue_card", "")
            prompts = question_info.get("prompts", [])
            if prompts:
                q_text += "\n" + "Prompts:\n" + "\n".join("- " + p for p in prompts)
        else:
            q_text = str(question_info)
    else:
        part_desc = "Part 3 - Two-way Discussion (abstract questions)"
        if isinstance(question_info, list):
            q_text = "\n".join(question_info)
        else:
            q_text = str(question_info)

    user_prompt = (
        "=== IELTS SPEAKING TEST ===\n"
        + "Part: " + part_desc + "\n\n"
        + "=== QUESTION ===\n"
        + q_text + "\n\n"
        + "=== STUDENT'S RESPONSE (for reference) ===\n"
        + (user_response if user_response else "[No response provided]") + "\n\n"
        + "=== TASK ===\n"
        + "Based on the question above, generate TWO distinct sample answers that would score Band 7.0 or higher.\n"
        + "Consider the student's vocabulary level and topic but demonstrate higher-level language.\n\n"
        + "=== FORMAT ===\n"
        + "---SAMPLE 1---\n"
        + "[Your first sample answer here]\n\n"
        + "---SAMPLE 2---\n"
        + "[Your second sample answer here]\n\n"
        + "IMPORTANT: Use exactly ---SAMPLE 1--- and ---SAMPLE 2--- as separators.\n"
        + "Do not include any other text outside of these two sample answers."
    )

    return system_prompt, user_prompt


def generate_sample_answers(part, question_info, user_response, topic=""):
    """Generate 2 Band 7+ sample answers using LLM API."""
    system_prompt, user_prompt = _build_prompt(part, question_info, user_response, topic)
    raw = _call_llm(system_prompt, user_prompt)
    if raw.startswith("[ERROR]"):
        return [raw, raw]
    samples = []
    for part_text in raw.split("---SAMPLE"):
        part_text = part_text.strip()
        if not part_text:
            continue
        if part_text[0].isdigit() and "---" in part_text[:5]:
            part_text = part_text[part_text.index("---") + 3:].strip()
        if len(part_text) > 20:
            samples.append(part_text)
    if len(samples) < 2:
        raw_lines = [l for l in raw.split("\n") if l.strip()]
        real = [l.strip() for l in raw_lines if len(l.strip()) > 50]
        if len(real) >= 2:
            samples = real[:2]
    while len(samples) < 2:
        samples.append("(Sample answer generation returned incomplete results. Please try again.)")
    return samples[:2]


def save_as_docx(sample_texts, filepath):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    title = doc.add_heading("IELTS Speaking - Reference Answer", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    for i, text in enumerate(sample_texts, 1):
        doc.add_heading("Sample Answer " + str(i), level=2)
        for pt in text.split(chr(10) + chr(10)):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(pt.strip())
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        doc.add_paragraph("")
    doc.add_paragraph("")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("-- Generated by IELTS Speaking Coach --")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.save(filepath)
    return filepath